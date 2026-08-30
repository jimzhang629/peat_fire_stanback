"""Minimal Markdown -> .docx converter for methods.md.

Handles: ATX headings, paragraphs, bullet/numbered lists, pipe tables,
fenced code blocks, and inline **bold** / *italic* / `code`. Display-math
blocks ($$...$$) are rendered as centred italic monospace, which is enough
for a Word draft a human will polish.
"""
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

INLINE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+?`|\$[^$]+?\$)")

# LaTeX -> unicode, enough to make the equations legible in Word without an
# equation editor. A human can drop in real Word equations later if wanted.
SYMBOLS = {
    r"\mathbb{E}": "E", r"\Sigma": "Σ", r"\sigma": "σ", r"\beta": "β",
    r"\theta": "θ", r"\mu": "μ", r"\Pr": "Pr", r"\log": "log",
    r"\times": "×", r"\le": "≤", r"\ge": "≥", r"\mid": " | ",
    r"\dots": "…", r"\top": "ᵀ", r"\perp": "⊥", r"\in": "∈",
    r"\sum": "Σ", r"\cdot": "·", r"\approx": "≈", r"\neq": "≠",
}
SUPERSCRIPT = {"-": "⁻", "1": "¹", "2": "²", "6": "⁶", "/": "ᐟ", "T": "ᵀ", "0": "⁰"}
SUBSCRIPT = {"0": "₀", "1": "₁", "2": "₂", "i": "ᵢ", "j": "ⱼ", "M": "M", "g": "g",
             "s": "ₛ", "t": "ₜ", "T": "T"}


def _strip_wrapper(tex, cmd):
    """Replace every ``\\cmd{body}`` with its brace-balanced ``body``."""
    out, needle = "", "\\" + cmd + "{"
    while True:
        k = tex.find(needle)
        if k < 0:
            return out + tex
        out += tex[:k]
        j, depth = k + len(needle), 1
        while j < len(tex) and depth:
            depth += (tex[j] == "{") - (tex[j] == "}")
            j += 1
        out += tex[k + len(needle):j - 1]
        tex = tex[j:]


def latex_to_text(tex):
    """Render a LaTeX fragment as readable plain text.

    Order matters: sizing/spacing macros go first (so ``\\left`` is not eaten by
    the ``\\le`` symbol rule), then accents, roots, scripts, and only then
    ``\\frac`` -- by which point its two arguments are brace-free.
    """
    # 1. font/formatting wrappers, sizing and spacing macros
    for cmd in ("mathrm", "mathbb", "mathbf", "text", "textit", "operatorname"):
        tex = _strip_wrapper(tex, cmd)
    tex = re.sub(r"\\(?:left|right|bigm|Bigg|bigg|Big|big|quad|qquad)\b", "", tex)
    tex = re.sub(r"\\[!,;:]", " ", tex).replace("\\ ", " ")

    # 2. accents and roots (innermost first, so scripts/frac see plain text)
    tex = re.sub(r"\\(?:hat|widehat)\{([^{}]*)\}", lambda m: m.group(1) + "\u0302", tex)
    tex = re.sub(r"\\hat\s*(\w)", lambda m: m.group(1) + "\u0302", tex)   # braceless \hat p
    tex = re.sub(r"\\(?:bar|overline)\{([^{}]*)\}", lambda m: m.group(1) + "\u0304", tex)

    # 3. sub/superscripts -> unicode where every character has a form
    def sup(m):
        body = m.group(1)
        return "".join(SUPERSCRIPT[c] for c in body) if all(c in SUPERSCRIPT for c in body) else f"^({body})"

    def sub(m):
        body = m.group(1)
        return "".join(SUBSCRIPT[c] for c in body) if all(c in SUBSCRIPT for c in body) else f"_{body}"

    tex = re.sub(r"\^\{([^{}]*)\}", sup, tex)
    tex = re.sub(r"_\{([^{}]*)\}", sub, tex)

    # 4. roots then fractions (their arguments are brace-free by now), then symbols
    for _ in range(3):
        tex = re.sub(r"\\sqrt\{([^{}]*)\}", r"√(\1)", tex)
    for _ in range(3):
        tex = re.sub(r"\\frac\{([^{}]*)\}\{([^{}]*)\}", r"(\1)/(\2)", tex)
    for k in sorted(SYMBOLS, key=len, reverse=True):
        tex = tex.replace(k, SYMBOLS[k])

    tex = re.sub(r"\\([A-Za-z]+)", r"\1", tex)          # any leftover command
    tex = tex.replace("{", "").replace("}", "")
    return re.sub(r"\s{2,}", " ", tex).strip()


def add_runs(par, text, bold=False):
    """Split `text` on inline markers and add styled runs to `par`.

    Recurses once into ``**bold**`` so code spans and math nested inside a bold
    phrase keep their own formatting instead of showing their markers.
    """
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            add_runs(par, tok[2:-2], bold=True)
            continue
        if tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        elif tok.startswith("$") and tok.endswith("$"):
            r = par.add_run(latex_to_text(tok[1:-1]))
            r.italic = True
            r.font.name = "Cambria Math"
        elif tok.startswith("*") and tok.endswith("*"):
            r = par.add_run(tok[1:-1])
            r.italic = True
        else:
            r = par.add_run(tok)
        r.bold = bold or r.bold


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_divider(line):
    return bool(re.fullmatch(r"\|[\s:\-|]+\|", line.strip()))


def convert(src, dest):
    lines = open(src, encoding="utf-8").read().split("\n")
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- horizontal rule ------------------------------------------------
        if stripped == "---":
            i += 1
            continue

        # --- blank ----------------------------------------------------------
        if not stripped:
            i += 1
            continue

        # --- fenced code / display math -------------------------------------
        if stripped.startswith("```"):
            i += 1
            body = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                body.append(lines[i])
                i += 1
            i += 1
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Pt(18)
            r = par.add_run("\n".join(body))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        if stripped.startswith("$$"):
            body = [stripped.strip("$").strip()]
            if not stripped.endswith("$$") or stripped == "$$":
                i += 1
                while i < len(lines) and not lines[i].strip().endswith("$$"):
                    body.append(lines[i].strip())
                    i += 1
                if i < len(lines):
                    body.append(lines[i].strip().rstrip("$").strip())
            i += 1
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = par.add_run(latex_to_text(" ".join(b for b in body if b)))
            r.italic = True
            r.font.name = "Cambria Math"
            r.font.size = Pt(11)
            continue

        # --- heading ---------------------------------------------------------
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(level=min(level, 4))
            add_runs(h, m.group(2))
            i += 1
            continue

        # --- table -----------------------------------------------------------
        if stripped.startswith("|") and i + 1 < len(lines) and is_divider(lines[i + 1]):
            header = split_row(stripped)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for cell, text in zip(table.rows[0].cells, header):
                cell.paragraphs[0].text = ""
                add_runs(cell.paragraphs[0], text)
            for row in rows:
                cells = table.add_row().cells
                for cell, text in zip(cells, row):
                    cell.paragraphs[0].text = ""
                    add_runs(cell.paragraphs[0], text)
            doc.add_paragraph()
            continue

        # --- list item (may wrap onto continuation lines) ---------------------
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            ordered = m.group(2)[0].isdigit()
            body = [m.group(3)]
            i += 1
            while i < len(lines):
                nxt = lines[i]
                if not nxt.strip() or re.match(r"^\s*([-*]|\d+\.)\s+", nxt) or nxt.strip().startswith(("#", "|", "```", "$$")):
                    break
                body.append(nxt.strip())
                i += 1
            style = "List Number" if ordered else "List Bullet"
            par = doc.add_paragraph(style=style)
            if indent >= 2:
                par.paragraph_format.left_indent = Pt(36)
            add_runs(par, " ".join(body))
            continue

        # --- paragraph --------------------------------------------------------
        body = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip() or re.match(r"^\s*([-*]|\d+\.)\s+", nxt) or nxt.strip().startswith(("#", "|", "```", "$$", "---")):
                break
            body.append(nxt.strip())
            i += 1
        par = doc.add_paragraph()
        add_runs(par, " ".join(body))

    doc.save(dest)
    print("wrote", dest)


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
