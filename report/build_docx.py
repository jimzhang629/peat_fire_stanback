"""Render a report markdown file into a Word document.

Defaults to project_report.md; pass another filename in report/ to render that
instead (e.g. report_draft.md). Placeholders ([[NUM: ...]], [[TODO: ...]], [[FIG n: ...]]) are
highlighted yellow, and the "> " guidance blocks are set in grey italic so they
read as scaffolding rather than as report prose.

    python report/build_docx.py [report_draft.md]

Re-run it after editing the markdown; it overwrites the .docx.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor

HERE = Path(__file__).resolve().parent
SRC = HERE / sys.argv[1] if len(sys.argv) > 1 else HERE / "project_report.md"
DST = SRC.with_suffix(".docx")

GREY = RGBColor(0x66, 0x66, 0x66)

# [[...]] placeholder, or **bold**, or *italic*, or `code`
INLINE = re.compile(r"(\[\[.*?\]\]|\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+`)")


def add_runs(paragraph, text, *, base_italic=False, base_color=None):
    """Write `text` into `paragraph`, honouring inline markdown and [[markers]]."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        run = paragraph.add_run()
        run.italic = base_italic
        if base_color is not None:
            run.font.color.rgb = base_color

        if piece.startswith("[[") and piece.endswith("]]"):
            run.text = piece
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.bold = True
        elif piece.startswith("**") and piece.endswith("**"):
            run.text = piece[2:-2]
            run.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            run.text = piece[1:-1]
            run.font.name = "Consolas"
        elif piece.startswith("*") and piece.endswith("*"):
            run.text = piece[1:-1]
            run.italic = True
        else:
            run.text = piece
    return paragraph


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator(line):
    return bool(re.fullmatch(r"\|[\s:\-|]+\|", line.strip()))


def main():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule -> page-ish separator
        if re.fullmatch(r"-{3,}", stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, "* * *", base_color=GREY)
            i += 1
            continue

        # Markdown table: header row, separator, body
        if stripped.startswith("|") and i + 1 < len(lines) and is_separator(lines[i + 1]):
            header = split_row(stripped)
            body = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                body.append(split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for cell, text in zip(table.rows[0].cells, header):
                cell.paragraphs[0].text = ""
                add_runs(cell.paragraphs[0], f"**{text}**")
            for row in body:
                cells = table.add_row().cells
                for cell, text in zip(cells, row):
                    cell.paragraphs[0].text = ""
                    add_runs(cell.paragraphs[0], text)
            doc.add_paragraph()
            continue

        # Headings
        m = re.match(r"(#{1,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading(level=min(level, 4))
            add_runs(heading, m.group(2))
            i += 1
            continue

        # Guidance blockquote -> grey italic, indented
        if stripped.startswith(">"):
            block = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(24)
            add_runs(p, " ".join(block).strip(), base_italic=True, base_color=GREY)
            continue

        # Bullets and numbered items
        m = re.match(r"[-*]\s+(.*)", stripped)
        if m:
            add_runs(doc.add_paragraph(style="List Bullet"), m.group(1))
            i += 1
            continue
        m = re.match(r"\d+\.\s+(.*)", stripped)
        if m:
            add_runs(doc.add_paragraph(style="List Number"), m.group(1))
            i += 1
            continue

        # Body paragraph: join wrapped lines until a blank or a new block
        block = []
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or re.match(r"(#{1,4}\s|[-*]\s|\d+\.\s|>|\|)", nxt) or re.fullmatch(r"-{3,}", nxt):
                break
            block.append(nxt)
            i += 1
        if not block:
            # The line opens a block the dispatch above did not claim (e.g. a
            # wrapped list item whose continuation starts with "|"). Emit it as
            # plain text; the guard is what keeps `i` moving.
            add_runs(doc.add_paragraph(), stripped)
            i += 1
            continue
        text = " ".join(block)
        # A whole-paragraph italic run is a figure/table caption.
        if text.startswith("*") and text.endswith("*") and not text.startswith("**"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(24)
            add_runs(p, text[1:-1], base_italic=True)
        else:
            add_runs(doc.add_paragraph(), text)

    doc.save(DST)
    print(f"wrote {DST}")


if __name__ == "__main__":
    main()
